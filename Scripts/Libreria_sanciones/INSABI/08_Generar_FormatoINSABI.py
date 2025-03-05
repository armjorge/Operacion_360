import os
import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import math
#from win32com.client import Dispatch
from docx import Document
from datetime import datetime
import locale
from PyPDF2 import PdfMerger
import zipfile
import platform
import sys
    


def open_excel(path):
    """Open an Excel file based on the operating system."""
    if not os.path.exists(path):
        print(f"Error: The file '{path}' does not exist.")
        return

    try:
        if platform.system() == "Windows":
            os.system(f'start excel "{path}"')
        elif platform.system() == "Darwin":  # macOS
            os.system(f'open -a "Microsoft Excel" "{path}"')
        else:
            print("Operating system not supported for opening Excel files.")
            return

        print(f"{os.path.basename(path)} opened successfully.")
    except Exception as e:
        print(f"Error trying to open the Excel file: {e}")

def open_docx(path):
    """Open a Word (.docx) file based on the operating system."""
    if not os.path.exists(path):
        print(f"Error: The file '{path}' does not exist.")
        return

    try:
        if platform.system() == "Windows":
            os.system(f'start winword "{path}"')
        elif platform.system() == "Darwin":  # macOS
            os.system(f'open -a "Microsoft Word" "{path}"')
        else:
            print("Operating system not supported for opening Word documents.")
            return

        print(f"{os.path.basename(path)} opened successfully.")
    except Exception as e:
        print(f"Error trying to open the Word document: {e}")

def open_pdf(pdf_path):
    """
    Opens the given PDF file in a system-compatible way.
    - Tries to use Adobe Acrobat if available.
    - Otherwise, uses the default PDF viewer.

    Parameters:
        pdf_path (str): The full path to the PDF file.

    Returns:
        None
    """
    if not os.path.exists(pdf_path):
        print(f"❌ Error: El archivo no existe: {pdf_path}")
        return

    system = platform.system()

    try:
        if system == "Windows":
            # Try opening with Acrobat Reader
            acrobat_path = r"C:\Program Files\Adobe\Acrobat DC\Acrobat\Acrobat.exe"
            if os.path.exists(acrobat_path):
                subprocess.run([acrobat_path, pdf_path], check=False)
            else:
                os.startfile(pdf_path)  # Open with default PDF viewer
        elif system == "Darwin":  # macOS
            acrobat_path = "/Applications/Adobe Acrobat DC/Adobe Acrobat.app"
            if os.path.exists(acrobat_path):
                subprocess.run(["open", "-a", acrobat_path, pdf_path], check=False)
            else:
                subprocess.run(["open", pdf_path], check=False)  # Default PDF viewer
        elif system == "Linux":
            subprocess.run(["xdg-open", pdf_path], check=False)

    except Exception as e:
        print(f"⚠️ No se pudo abrir el archivo: {e}")

def A_confirm_files_exists(folder_name): 
    print("Looking for the folder...")
    
    # Check if the folder exists
    folder_path = os.path.join('.', folder_name)
    if not os.path.exists(folder_path):
        print(f"Folder '{folder_name}' not found, try again")
        return None, None
    
    print("Folder found")
    print("Trying to find the Audit and Relacion Excel files...")
    
    # Define paths for the audit and relacion files
    audit_xlsx = os.path.join(folder_path, f"{folder_name}_audit.xlsx")
    relacion_xlsx = os.path.join(folder_path, f"{folder_name}_relacion.xlsx")
    
    # Check if both files exist
    missing_files = []
    if not os.path.exists(audit_xlsx):
        missing_files.append(f"{folder_name}_audit.xlsx")
    if not os.path.exists(relacion_xlsx):
        missing_files.append(f"{folder_name}_relacion.xlsx")
    
    if missing_files:
        print(f"*****************\nThe following file(s) are missing: {', '.join(missing_files)}")
        return None, None
    
    print("Audit and Relacion files present")
    print("Loading local dataframes...")
    
    # Load the Excel files into dataframes
    try:
        df_audit = pd.read_excel(audit_xlsx)
        df_relacion = pd.read_excel(relacion_xlsx)
    except Exception as e:
        print("Error: Failed to load Excel files. Exception:", str(e))
        return None, None

    # Diagnose if the required column is missing before attempting to process it
    if 'ORDEN DE SUMINISTRO' not in df_audit.columns:
        print("Error: The column 'ORDEN DE SUMINISTRO' is missing in the Audit Excel file.")
        print("Please verify that the file contains this column and that its name is correct.")
        return None, None

    if 'ORDEN DE SUMINISTRO' not in df_relacion.columns:
        print("Error: The column 'ORDEN DE SUMINISTRO' is missing in the Relacion Excel file.")
        print("Please verify that the file contains this column and that its name is correct.")
        return None, None

    # Remove any whitespace (spaces, tabs, newlines, etc.) from the 'ORDEN DE SUMINISTRO' column
    df_audit['ORDEN DE SUMINISTRO'] = df_audit['ORDEN DE SUMINISTRO'].astype(str).str.replace(r'\s', '', regex=True)
    df_relacion['ORDEN DE SUMINISTRO'] = df_relacion['ORDEN DE SUMINISTRO'].astype(str).str.replace(r'\s', '', regex=True)
    
    print("Dataframes loaded, ready for the next step")
    return df_audit, df_relacion

def B_carga_desde_el_GOOGLESHEET(sheets_in_management):
    # Constants for Google Sheets API
    spreadsheet_url = 'https://docs.google.com/spreadsheets/d/1KN4XwXQlZ5jhyErxdpA_R6kNk2tKwwFeDnCDuHCr0fo/edit#gid=2033397596'
    credentials_file = 'key.json'
    
    # Setup Google Sheets API
    scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
    creds = ServiceAccountCredentials.from_json_keyfile_name(credentials_file, scope)
    client = gspread.authorize(creds)
    
    # Open the spreadsheet by URL
    spreadsheet = client.open_by_url(spreadsheet_url)
    
    # Dictionary to store the resulting dataframes
    dfs = {}
    
    # Function to load data from a worksheet into a DataFrame
    def worksheet_to_df(spreadsheet, worksheet_name, columns_range):
        worksheet = spreadsheet.worksheet(worksheet_name)
        data = worksheet.get(columns_range)
        return pd.DataFrame(data[1:], columns=data[0])

    # Iterate over the sheets_in_management dictionary and create dataframes
    print("Loading from google sheet")
    for df_name, (sheet_name, columns_range) in sheets_in_management.items():
        try:
            df = worksheet_to_df(spreadsheet, sheet_name, columns_range)
            dfs[df_name] = df
        except Exception as e:
            print(f"An error occurred while loading {sheet_name}: {e}")
            dfs[df_name] = pd.DataFrame()  # Return an empty DataFrame in case of an error
    print("Dataframes loaded, returned and ready")
    return dfs

def C_generating_single_dataframe(df_INSABI_gsheet, df_audit, df_relacion):
    global folder_name
    print("Checking for the needed columns...")

    # Define the required columns for each dataframe
    required_columns_INSABI = ['Contrato', 'Factura', 'UUID', 'NÚMERO DE ORDEN DE SUMINISTRO', 'Importe']
    required_columns_audit = ['ORDEN DE SUMINISTRO', 'Filename', 'XML_UUID_NC', 'XML_Total']
    required_columns_relacion = ['ORDEN DE SUMINISTRO', 'PENA', 'Folio NC']

    # Check for missing columns in each dataframe
    missing_columns = {
        'df_INSABI_gsheet': [col for col in required_columns_INSABI if col not in df_INSABI_gsheet.columns],
        'df_audit': [col for col in required_columns_audit if col not in df_audit.columns],
        'df_relacion': [col for col in required_columns_relacion if col not in df_relacion.columns],
    }

    # Collect all missing columns for user feedback
    missing_feedback = []
    for df_name, cols in missing_columns.items():
        if cols:  # If there are missing columns
            missing_feedback.append(f"*****************\n{df_name} is missing columns: {', '.join(cols)}\n*****************\n")

    # If any columns are missing, notify the user and exit the function
    if missing_feedback:
        print("The following columns are missing:")
        for feedback in missing_feedback:
            print(feedback)
        print("Please check the dataframes and try again.")
        return None  # Exit function if columns are missing

    print("Columns found, proceeding")

    # Create the master dataframe with the required headers
    headers = [
        "Columna1", "PROVEEDOR", "CONTRATO", "NUMERO DE FACTURA", "FOLIO FISCAL FACTURA",
        "ORDEN DE SUMINISTRO", "MONTO FACTURADO", "NOTA DE CREDITO (FOLIO Y SERIE)",
        "FOLIO FISCAL NOTA DE CREDITO", "MONTO SANCIONADO", "NUMERO DE CEDULA"
    ]
    master_df = pd.DataFrame(columns=headers)

    # Populate the master dataframe
    master_df['Columna1'] = range(1, len(df_relacion) + 1)  # Consecutive number of rows
    master_df['PROVEEDOR'] = "ESEOTRES PHARMA"  # Fixed value
    master_df['ORDEN DE SUMINISTRO'] = df_relacion['ORDEN DE SUMINISTRO']  # Direct mapping
    master_df['CONTRATO'] = master_df['ORDEN DE SUMINISTRO'].apply(
        lambda x: C1_returning_values(df_INSABI_gsheet, x, 'NÚMERO DE ORDEN DE SUMINISTRO', 'Contrato')
    )
    master_df['NUMERO DE FACTURA'] = master_df['ORDEN DE SUMINISTRO'].apply(
        lambda x: C1_returning_values(df_INSABI_gsheet, x, 'NÚMERO DE ORDEN DE SUMINISTRO', 'Factura')
    )
    master_df['FOLIO FISCAL FACTURA'] = master_df['ORDEN DE SUMINISTRO'].apply(
        lambda x: C1_returning_values(df_INSABI_gsheet, x, 'NÚMERO DE ORDEN DE SUMINISTRO', 'UUID')
    )
    master_df['MONTO FACTURADO'] = master_df['ORDEN DE SUMINISTRO'].apply(
        lambda x: C1_returning_values(df_INSABI_gsheet, x, 'NÚMERO DE ORDEN DE SUMINISTRO', 'Importe')
    )
    master_df['NOTA DE CREDITO (FOLIO Y SERIE)'] = master_df['ORDEN DE SUMINISTRO'].apply(
        lambda x: C1_returning_values(df_relacion, x, 'ORDEN DE SUMINISTRO', 'Folio NC')
    )
    master_df['FOLIO FISCAL NOTA DE CREDITO'] = master_df['NOTA DE CREDITO (FOLIO Y SERIE)'].apply(
        lambda x: C1_returning_values(df_audit, x, 'Filename', 'XML_UUID_NC')
    )
    master_df['MONTO SANCIONADO'] = master_df['ORDEN DE SUMINISTRO'].apply(
        lambda x: C1_returning_values(df_relacion, x, 'ORDEN DE SUMINISTRO', 'PENA')
    )
    master_df['NUMERO DE CEDULA'] = folder_name  # Fixed value from global variable

    return master_df

def C1_returning_values(dataframe, origin_column, match_column, return_value):
    """
    Helper function to find and return a value from a dataframe based on a match.
    """
    return dataframe.loc[dataframe[match_column] == origin_column, return_value].values[0] if not dataframe.loc[dataframe[match_column] == origin_column].empty else None

""" Función D_convert_formato_INSABI_toPDF para windows solamente
def D_convert_formato_INSABI_toPDF(excel_file, orientation):
    
    #Convert an Excel file to a PDF with specified orientation (landscape or portrait).
    
    #:param excel_file: Path to the Excel file.
    #:param orientation: Either 'landscape' or 'portrait'.
    
    global folder_name
    pdf_file = os.path.abspath(f"./{folder_name}/{folder_name}_formato_INSABI.pdf")

    # Define orientation values
    orientation_map = {
        'landscape': 2,  # xlLandscape
        'portrait': 1    # xlPortrait
    }

    if orientation not in orientation_map:
        raise ValueError("Invalid orientation. Use 'landscape' or 'portrait'.")

    try:
        # Initialize Excel application
        excel_app = Dispatch("Excel.Application")
        excel_app.Visible = False  # Keep Excel in the background

        # Open the Excel file
        workbook = excel_app.Workbooks.Open(os.path.abspath(excel_file))

        # Adjust settings for each sheet in the workbook
        for sheet in workbook.Sheets:
            sheet.PageSetup.Orientation = orientation_map[orientation]
            sheet.PageSetup.PaperSize = 1  # xlPaperLetter
            sheet.PageSetup.Zoom = False  # Disable zoom
            sheet.PageSetup.FitToPagesWide = 1
            sheet.PageSetup.FitToPagesTall = False  # Fit all columns to one page

        # Export to PDF
        workbook.ExportAsFixedFormat(
            Type=0,  # xlTypePDF
            Filename=pdf_file,
            Quality=0,  # xlQualityStandard
            IncludeDocProperties=True,
            IgnorePrintAreas=False,
            OpenAfterPublish=False
        )

        workbook.Close(SaveChanges=False)
        excel_app.Quit()
        print(f"Excel file successfully converted to PDF and saved at: {pdf_file}")

    except Exception as e:
        print(f"An error occurred: {e}")
        if 'excel_app' in locals():
            excel_app.Quit()
        raise
""" 


def D_convert_formato_INSABI_toPDF(excel_file, orientation):
    """
    Convert an Excel file to a PDF with specified orientation (landscape or portrait).

    :param excel_file: Path to the Excel file.
    :param orientation: Either 'landscape' or 'portrait'.
    """
    global folder_name
    pdf_file = os.path.abspath(f"./{folder_name}/{folder_name}_formato_INSABI.pdf")

    orientation_map = {
        'landscape': 2,  # xlLandscape
        'portrait': 1    # xlPortrait
    }

    if orientation not in orientation_map:
        raise ValueError("Invalid orientation. Use 'landscape' or 'portrait'.")

    system = platform.system()

    try:
        if system == "Windows":
            from win32com.client import Dispatch

            # Initialize Excel application
            excel_app = Dispatch("Excel.Application")
            excel_app.Visible = False

            # Open the Excel file
            workbook = excel_app.Workbooks.Open(os.path.abspath(excel_file))

            # Adjust settings
            for sheet in workbook.Sheets:
                sheet.PageSetup.Orientation = orientation_map[orientation]
                sheet.PageSetup.PaperSize = 1
                sheet.PageSetup.Zoom = False
                sheet.PageSetup.FitToPagesWide = 1
                sheet.PageSetup.FitToPagesTall = False

            # Export to PDF
            workbook.ExportAsFixedFormat(
                Type=0,
                Filename=pdf_file,
                Quality=0,
                IncludeDocProperties=True,
                IgnorePrintAreas=False,
                OpenAfterPublish=False
            )

            workbook.Close(SaveChanges=False)
            excel_app.Quit()

        elif system == "Darwin":  # macOS
            print(f"\n*********************\n, La funcionalidad aún no está \n Guarda a mano el excel que se abrió en la carpeta {os.path.basename(folder_name)}")
            
            from appscript import app, k
            import mactypes

            # Wrap the Excel file path as a file reference
            excel_app = app('Microsoft Excel')
            workbook = excel_app.open(mactypes.File(os.path.abspath(excel_file)))
            
            if workbook is None:
                raise Exception("Failed to open workbook. Please verify the file path and Excel configuration.")

            # Iterate over each worksheet and set up page properties
            for sheet in workbook.worksheets():
                sheet.page_setup.orientation.set(k.landscape if orientation == 'landscape' else k.portrait)
                sheet.page_setup.paper_size.set(k.paper_letter)
                sheet.page_setup.zoom.set(False)
                sheet.page_setup.fit_to_pages_wide.set(1)
                sheet.page_setup.fit_to_pages_tall.set(0)

            # Wrap the output PDF file path too
            workbook.save_as(mactypes.File(os.path.abspath(pdf_file)), file_format=k.PDF)
            workbook.close()

        else:
            print("Unsupported operating system for PDF conversion.")
            return

        print(f"PDF created successfully: {pdf_file}")

    except Exception as e:
        print(f"An error occurred: {e}")

def E_generateletter(input_dataframe):
    """
    Generate a letter document by replacing placeholders in a Word template,
    then convert it to a PDF using the appropriate method for the OS.
    
    :param input_dataframe: DataFrame containing the data for the letter.
    """
    import os
    import locale
    from datetime import datetime
    from docx import Document
    import platform

    global folder_name
    print(f"\nThe function generate letter will generate {folder_name}_carta.pdf")
    
    # Use current working directory as the root folder
    root_dir = os.getcwd()
    
    # Build absolute paths based on the root folder
    doc_template = os.path.join(root_dir, "Templates", "Carátula.docx")
    output_doc = os.path.join(root_dir, folder_name, f"{folder_name}_carta.docx")
    output_pdf = os.path.join(root_dir, folder_name, f"{folder_name}_carta.pdf")
    
    # Confirm that the template file exists
    if not os.path.exists(doc_template):
        print(f"Error: Template file not found at '{doc_template}'")
        return
    
    # Set locale to Spanish for correct month names
    try:
        locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
    except locale.Error:
        print("Warning: Spanish locale not available. Using default locale.")
    
    # Ask for the issue date
    today = datetime.today()
    user_date = input("Use today's date? (yes or enter date in format dd/mm): ").strip()
    if user_date.lower() == "yes":
        issue_date = today.strftime("%d/%m")
    else:
        issue_date = user_date

    try:
        day, month = map(int, issue_date.split("/"))
        date_str = f"Ciudad de México a {day} de {today.strftime('%B')} del {today.year}"
        date_str = date_str.replace(today.strftime('%B'), today.strftime('%B').capitalize())
    except ValueError:
        print("Invalid date format. Please use dd/mm.")
        return

    # Count rows from the DataFrame
    rows_number = len(input_dataframe)
    fields_to_replace = {
        'date': date_str,
        'folder_name': folder_name,
        'rows_number': rows_number
    }
    
    # Load the template document
    try:
        doc = Document(doc_template)
    except Exception as e:
        print(f"An error occurred while loading the template: {e}")
        return

    # Replace placeholders in all paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in fields_to_replace.items():
            if f"{{{placeholder}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{placeholder}}}", str(value))
    
    # Populate table data if a table exists in the document
    try:
        word_table = 1  # Using a 1-based index for the target table
        word_headers_row = 1  # 1-based index for header row
        if len(doc.tables) < word_table:
            print("Warning: No table found in the template. Skipping table population.")
        else:
            table = doc.tables[word_table - 1]
            headers_row = word_headers_row - 1
            word_headers = [cell.text.strip() for cell in table.rows[headers_row].cells]
            # Append a new row for each row in the DataFrame
            for _, row_data in input_dataframe.iterrows():
                new_row = table.add_row()
                for idx, header in enumerate(word_headers):
                    if header in input_dataframe.columns:
                        new_row.cells[idx].text = str(row_data[header])
    except Exception as e:
        print(f"An error occurred during table population: {e}")
    
    # Save the updated document
    try:
        doc.save(output_doc)
        print(f"Letter successfully generated and saved at {output_doc}")
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")
        return

    # Open the document for manual formatting using your cross-platform function
    try:
        open_docx(output_doc)
    except Exception as e:
        print(f"An error occurred while opening the document: {e}")
    
    input("Fix the font and format, then press Enter to continue.")
    
    # Convert the Word document to PDF using OS-specific methods
    try:
        system = platform.system()
        if system == "Windows":
            from win32com.client import Dispatch
            word_app = Dispatch("Word.Application")
            word_app.Visible = False
            word_doc = word_app.Documents.Open(output_doc)
            # FileFormat=17 corresponds to PDF
            word_doc.SaveAs(output_pdf, FileFormat=17)
            word_doc.Close()
            word_app.Quit()
        elif system == "Darwin":  # macOS
            from appscript import app, k
            word_app = app('Microsoft Word')
            # Open the document via the documents collection
            word_doc = word_app.documents.open(output_doc)
            # Save as PDF using the PDF file format constant
            word_doc.save_as(os.path.abspath(output_pdf), file_format=k.pdf)
            word_doc.close()
        else:
            print("Unsupported operating system for Word to PDF conversion.")
            return

        print(f"PDF successfully generated and saved at {output_pdf}")
    except Exception as e:
        print(f"An error occurred during PDF conversion: {e}")
        try:
            if 'word_app' in locals():
                if system == "Windows":
                    word_app.Quit()
                elif system == "Darwin":
                    word_doc.close()
        except Exception:
            pass

def F_estatusarmado(folder_name):
    """
    Check the status of required files, merge them into a single PDF if all exist, 
    and save the merged file as 'Portada'.
    """
    

    # Define file paths
    print("\nA function will try to merge relation, excel and portada\n")
    Cedula = os.path.abspath(f"./{folder_name}/{folder_name}.pdf")
    Formato_excel_INSABI = os.path.abspath(f"./{folder_name}/{folder_name}_formato_INSABI.pdf")
    Carta = os.path.abspath(f"./{folder_name}/{folder_name}_carta.pdf")
    Portada = os.path.abspath(f"./{folder_name}/{folder_name}_portada.pdf")

    # List of required files
    required_files = {
        "Cedula": Cedula,
        "Formato_excel_INSABI": Formato_excel_INSABI,
        "Carta": Carta
    }

    # Check for missing files
    missing_files = [name for name, path in required_files.items() if not os.path.exists(path)]
    if missing_files:
        print("The following files are missing:")
        for name in missing_files:
            print(f"- {name}")
        return  # Exit if any files are missing

    # Merge files if all exist
    try:
        merger = PdfMerger()
        for file_path in required_files.values():
            merger.append(file_path)
        merger.write(Portada)
        merger.close()
        print(f"Merged file saved as {Portada}")

    except Exception as e:
        print(f"An error occurred while merging files: {e}")
        return

    print("Ready to merge Portada with NC's files. Go to the next step.")

def G_generatezip(folder_name):
    """
    #Generate a ZIP file containing the `Formato_INSABI.xlsx` and all files starting with 'NC'.
    """

    # Define paths
    base_dir = os.path.abspath(f"./{folder_name}")
    zip_path = os.path.join(base_dir, f"{folder_name}.zip")
    formato_file = os.path.join(base_dir, f"{folder_name}_formato_INSABI.xlsx")
    
    # Collect all files to include in the ZIP
    files_to_zip = []
    if os.path.exists(formato_file):
        files_to_zip.append(formato_file)
    else:
        print(f"Missing file: {formato_file}")

    # Include all files starting with 'NC'
    for file in os.listdir(base_dir):
        if file.startswith("NC") and os.path.isfile(os.path.join(base_dir, file)):
            files_to_zip.append(os.path.join(base_dir, file))

    # Check if there are files to zip
    if not files_to_zip:
        print("No files found to include in the ZIP.")
        return

    # Create ZIP file
    try:
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in files_to_zip:
                zipf.write(file, os.path.basename(file))
        print(f"ZIP with NC files + Formato_INSABI was created: {zip_path}")
    except Exception as e:
        print(f"An error occurred while creating the ZIP file: {e}")
        
folder_name = input("Enter the folder name: ") 

def main(): 
    global folder_name
    df_audit, df_relacion = A_confirm_files_exists(folder_name)
    
    if df_audit is not None and df_relacion is not None:
        # Proceed with further processing
        print("Dataframes are ready for merging or further processing.")
    else:
        print("File check failed. Please try again.")

    ### Cargar el RAW_INSABI ###

    cargar_desde_GSHEET = {
        'df_RAW_INSABI': ('RAW_INSABI', 'A:AA')#,'df_anotherSheet': ('Another_Sheet_Name', 'A:B')
        }
    dataframes = B_carga_desde_el_GOOGLESHEET(cargar_desde_GSHEET)    
    df_RAW_INSABI = dataframes['df_RAW_INSABI']
    if df_RAW_INSABI is not None:
        # Proceed with further processing
        print("")
    else:
        print("File check failed. Please try again.")    
        
    master_df = C_generating_single_dataframe(df_RAW_INSABI, df_audit,df_relacion)
    
    if master_df is not None:
        print("Dataframe generation complete. Apply format before converting it to PDF")
        output_file = f"./{folder_name}/{folder_name}_formato_INSABI.xlsx"
        template_file = "./Templates/Template.xlsx"        
        # Save `master_df` to the output file
        excel_rewritting = input("We have a dataframe with the needed data. Do we rewrite the Excel file? (Yes or No): ").strip().lower()
        if excel_rewritting == "yes":
            master_df.to_excel(output_file, index=False)
            print(f"Excel file saved to {output_file}.")
        else:
            print("Understood. Skipping Excel file rewrite.\nApply format, the excel files will be opened\n")
    else:
        print("Dataframe generation failed. Please address the missing columns and try again.")
        
    template_file = os.path.abspath(r"./Templates/Template.xlsx")  # Ensure absolute path
    output_file = os.path.abspath(f"./{folder_name}/{folder_name}_formato_INSABI.xlsx")  # Ensure absolute path
    # Check existence before opening
    if os.path.exists(template_file):
        open_excel(template_file)
    else:
        print(f"Template file not found: {template_file}")

    if os.path.exists(output_file):
        open_excel(output_file)
    else:
        print(f"Output file not found: {output_file}")
        
    input("Push enter when you finish applying the format and want to continue: ")
    D_convert_formato_INSABI_toPDF(output_file, 'landscape')    
    E_generateletter(master_df)
    G_generatezip(folder_name)
    F_estatusarmado(folder_name)

    print("08_Generar_FormatoINSABI.py finished, please run 09_PDFprint.py")


if __name__ == "__main__":
    main()
