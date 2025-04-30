import pandas as pd
from datetime import datetime
from docx import Document
from docx2pdf import convert
import os



def get_spanish_date(date_str):
    months = {
        '01': 'Enero', '02': 'Febrero', '03': 'Marzo', '04': 'Abril',
        '05': 'Mayo', '06': 'Junio', '07': 'Julio', '08': 'Agosto',
        '09': 'Septiembre', '10': 'Octubre', '11': 'Noviembre', '12': 'Diciembre'
    }
    day, month, year = date_str.split('/')
    return f"{day} de {months[month]} de {year}"

def replace_text(doc, context):
    for p in doc.paragraphs:
        for key, value in context.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(f'{{{key}}}', value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, context)

def generate_letter(template_path, output_path, context):
    doc = Document(template_path)
    replace_text(doc, context)
    doc.save(output_path)

def letter_simple():
    
    # Load the data from the Excel file
    df_clientes = pd.read_excel('./Bases/db_clientes_personas_testigos.xlsx', sheet_name='Clientes')
    df_reciben_poder = pd.read_excel('./Bases/db_clientes_personas_testigos.xlsx', sheet_name='Equipo')
    df_testigos = pd.read_excel('./Bases/db_clientes_personas_testigos.xlsx', sheet_name='Testigos')
    today = datetime.today().strftime("%d/%m/%Y")
    word_sin_testigos = './Cartas poder/plantilla_sintestigos.docx'
    word_con_testigos = './Cartas poder/plantilla_contestigos.docx'
    
    # 1 Define date
    use_today = input("Use today date? (yes/no) ")
    if use_today.lower() == 'yes':
        date = today
    else:
        valid_date = False
        while not valid_date:
            try:
                date_input = input("Please input the desired date (dd/mm/yyyy): ")
                datetime.strptime(date_input, "%d/%m/%Y")
                date = date_input
                valid_date = True
            except ValueError:
                print("Invalid date format. Please try again.")

    date_str = get_spanish_date(date)

    # 2 Define client
    print(df_clientes[['Client']])
    client_index = int(input("We'll issue the letter to the client number? "))
    client = df_clientes.at[client_index, 'Client']

    # 3 Define team member
    print(df_reciben_poder[['granted']])
    granted_index = int(input("Who'll receive the power? "))
    granted = df_reciben_poder.at[granted_index, 'granted']

    context = {
        'date': date_str,
        'client': client,
        'granted': granted
    }
    
    output_path = f'./{client} {granted}.docx'

    # 4 Generate the PDF 
    generate_letter(word_sin_testigos, output_path, context)
    pdf_output_path = f'./{client} {granted}.pdf'
    convert(output_path)
    os.rename(output_path.replace('.docx', '.pdf'), pdf_output_path)
    os.remove(output_path)
    
    print(f"Letter generated for {client} granting {granted} for the day {date_str}")